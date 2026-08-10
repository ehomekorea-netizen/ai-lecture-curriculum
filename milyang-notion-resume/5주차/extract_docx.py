import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('디지털 실무 인터랙티브 웹 포트폴리오 5주차 가이드.docx')

with open('docx_extracted.txt', 'w', encoding='utf-8') as f:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f'[{i}] STYLE={para.style.name} | {para.text}\n')

    # Also extract tables
    f.write('\n\n=== TABLES ===\n')
    for ti, table in enumerate(doc.tables):
        f.write(f'\n--- TABLE {ti} ---\n')
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            f.write(' | '.join(cells) + '\n')

print('Done')
