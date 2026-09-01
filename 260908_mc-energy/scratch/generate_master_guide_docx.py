import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_master_guide():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Palette
    PRIMARY = RGBColor(14, 76, 146)     # Deep Navy Blue (#0E4C92)
    SECONDARY = RGBColor(2, 132, 199)   # Vibrant Sky Blue (#0284C7)
    ACCENT = RGBColor(16, 185, 129)     # Emerald Green (#10B981)
    DARK_TEXT = RGBColor(30, 41, 59)    # Slate 800 (#1E293B)
    MUTED_TEXT = RGBColor(100, 116, 139) # Slate 500 (#64748B)

    # Helper: Set Cell Background Color
    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Helper: Set Cell Padding
    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Helper: Set Cell Border
    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = parse_xml(f'<{tag} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" w:sz="{edge_data.get("sz", "4")}" w:space="0" w:color="{edge_data.get("color", "auto")}"/>')
                tcBorders.append(element)
            else:
                tag = 'w:{}'.format(edge)
                element = parse_xml(f'<{tag} {nsdecls("w")} w:val="none"/>')
                tcBorders.append(element)
        tcPr.append(tcBorders)

    # Header / Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_tag = title_p.add_run("LECTURE PRODUCTION BLUEPRINT")
    run_tag.font.name = "Pretendard"
    run_tag.font.size = Pt(9.5)
    run_tag.font.bold = True
    run_tag.font.color.rgb = SECONDARY

    h1_p = doc.add_paragraph()
    h1_p.paragraph_format.space_before = Pt(0)
    h1_p.paragraph_format.space_after = Pt(6)
    run_h1 = h1_p.add_run("Talks-Grade Slidev 고품격 강의 덱 제작 마스터 로드맵")
    run_h1.font.name = "Pretendard"
    run_h1.font.size = Pt(22)
    run_h1.font.bold = True
    run_h1.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("기획(Lecture Director)부터 에셋 맥락 매핑, 120 FPS 슬라이드 프로덕션(Talks Slidev), 그리고 6대 무결점 품질 가드레일까지")
    run_sub.font.name = "Pretendard"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = MUTED_TEXT

    # Callout Box: Core Philosophy
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_shading(cell, "F0F9FF")
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    set_cell_border(cell, left={"val": "single", "sz": "24", "color": "0284C7"})
    
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)
    c_run1 = cp.add_run("💡 마스터 워크플로우 한 줄 요약\n")
    c_run1.font.name = "Pretendard"
    c_run1.font.bold = True
    c_run1.font.size = Pt(11)
    c_run1.font.color.rgb = PRIMARY
    c_run2 = cp.add_run("클라이언트 계획서/영상 접수 ➔ ① lecture-director로 맥락&브리프 확정 ➔ ② 에셋+URL 맥락 매핑 ➔ ③ talks-slidev로 120 FPS 덱 완성 ➔ ④ 6대 무결점 QA & 로컬 커밋")
    c_run2.font.name = "Pretendard"
    c_run2.font.size = Pt(10)
    c_run2.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─────────────────────────────────────────────────────────
    # SECTION 1: 4-Phase Architecture
    # ─────────────────────────────────────────────────────────
    sec1 = doc.add_paragraph()
    sec1_run = sec1.add_run("1. 강의 덱 제작 4단계 엔드투엔드 파이프라인")
    sec1_run.font.name = "Pretendard"
    sec1_run.font.size = Pt(15)
    sec1_run.font.bold = True
    sec1_run.font.color.rgb = PRIMARY
    sec1.paragraph_format.space_before = Pt(14)
    sec1.paragraph_format.space_after = Pt(8)

    # Table for 4 Phases
    p_table = doc.add_table(rows=5, cols=4)
    p_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_table.autofit = False
    col_widths = [Inches(1.1), Inches(1.8), Inches(2.2), Inches(1.7)]
    
    headers = ["단계 (Phase)", "주요 활동 & 도구", "핵심 산출물 / 행동", "담당 스킬"]
    for i, h in enumerate(headers):
        c = p_table.cell(0, i)
        c.width = col_widths[i]
        set_cell_shading(c, "0E4C92")
        set_cell_margins(c, 100, 100, 120, 120)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Pretendard"
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    phase_data = [
        ("Phase 1\n기획 & 브리프", "클라이언트 의뢰서 분석\nGrill-Me 역질문 인터뷰\n최신 6개월 AI 트렌드 검증", "차시별 설계도 파일\n(sessions/S##-brief.md)\n필요 에셋 리스트 처방", "lecture-director"),
        ("Phase 2\n에셋 & 맥락 매핑", "킬러 에셋(이미지/영상) 수집\n기사 원문 링크(URL) 매핑\npublic/ 폴더에 에셋 적재", "URL + 파일명 1줄 전달\n(원문 팩트/수치 추출 준비)", "강사 ➔ 에이전트\n(URL 크롤링 연동)"),
        ("Phase 3\n덱 프로덕션", "Talks 보일러플레이트 세팅\n6대 표준 슬라이드 렌더링\nAnti-AI 발표 대본 자동 작성", "전체 Slidev 슬라이드덱\n(slides.md & Vue 컴포넌트)\n수강생 프롬프트 치트키", "talks-slidev"),
        ("Phase 4\n무결점 QA & 커밋", "6대 무결점 가드레일 전수 검수\npnpm run build 무결점 통과\n원격 Push 금지 로컬 커밋", "120 FPS 무결점 덱 완결\nQA_REPORT.md 업데이트\n안전한 Git 로컬 커밋", "talks-slidev\n(QA 프로토콜)")
    ]

    for row_idx, data in enumerate(phase_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            c = p_table.cell(row_idx, col_idx)
            c.width = col_widths[col_idx]
            set_cell_shading(c, bg)
            set_cell_margins(c, 100, 100, 120, 120)
            set_cell_border(c, bottom={"val": "single", "sz": "4", "color": "E2E8F0"})
            p = c.paragraphs[0]
            if col_idx in (0, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = "Pretendard"
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = PRIMARY
            elif col_idx == 3:
                r.font.bold = True
                r.font.color.rgb = SECONDARY
            else:
                r.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─────────────────────────────────────────────────────────
    # SECTION 2: Phase 1 Deep-Dive (lecture-director)
    # ─────────────────────────────────────────────────────────
    sec2 = doc.add_paragraph()
    sec2_run = sec2.add_run("2. Phase 1: 기획 & 브리프 확정 (`lecture-director`)")
    sec2_run.font.name = "Pretendard"
    sec2_run.font.size = Pt(14)
    sec2_run.font.bold = True
    sec2_run.font.color.rgb = PRIMARY
    sec2.paragraph_format.space_before = Pt(14)
    sec2.paragraph_format.space_after = Pt(6)

    def add_bullet_point(p, title, desc):
        r1 = p.add_run(f"• {title}: ")
        r1.font.name = "Pretendard"
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = PRIMARY
        r2 = p.add_run(desc + "\n")
        r2.font.name = "Pretendard"
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK_TEXT

    b_p = doc.add_paragraph()
    b_p.paragraph_format.space_after = Pt(6)
    add_bullet_point(b_p, "손에 있는 만큼 먼저 던지기", "완벽하게 에셋을 다 모을 필요 없이 계획서(docx/pdf)와 참고 유튜브 링크만 먼저 주면 디렉터가 뼈대를 잡고 필요한 에셋을 역으로 처방합니다.")
    add_bullet_point(b_p, "Grill-Me 인터뷰로 강사 의도 동기화", "이번 차시의 [개념 vs 실습] 시간 배분(예: 20/30, 15/35), 수강생 수준, 수료 후 남길 단 하나의 핵심 가치(One Thing)를 확정합니다.")
    add_bullet_point(b_p, "최신 6개월 AI 트렌드 보장", "과거의 레거시 프롬프트 꼼수를 배제하고 Thinking 자가보정 추론, Work 도구 실행, RAG 팩트 검증 등 최신 SOTA 기술을 자동 반영합니다.")

    # ─────────────────────────────────────────────────────────
    # SECTION 3: Phase 2 Deep-Dive (Asset & URL Mapping)
    # ─────────────────────────────────────────────────────────
    sec3 = doc.add_paragraph()
    sec3_run = sec3.add_run("3. Phase 2: 킬러 에셋 수집 & URL 맥락 매핑 3대 패턴")
    sec3_run.font.name = "Pretendard"
    sec3_run.font.size = Pt(14)
    sec3_run.font.bold = True
    sec3_run.font.color.rgb = PRIMARY
    sec3.paragraph_format.space_before = Pt(14)
    sec3.paragraph_format.space_after = Pt(6)

    # Callout for Pattern 1
    t_pat = doc.add_table(rows=1, cols=1)
    t_pat.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_pat = t_pat.cell(0, 0)
    c_pat.width = Inches(6.8)
    set_cell_shading(c_pat, "F0FDF4")
    set_cell_margins(c_pat, 120, 120, 160, 160)
    set_cell_border(c_pat, left={"val": "single", "sz": "20", "color": "10B981"})
    
    pat_p = c_pat.paragraphs[0]
    pat_p.paragraph_format.space_after = Pt(0)
    pr1 = pat_p.add_run("🌟 가장 추천하는 실전 에셋 전달법 (URL + 파일명 1줄 매핑)\n")
    pr1.font.name = "Pretendard"
    pr1.font.bold = True
    pr1.font.size = Pt(10.5)
    pr1.font.color.rgb = ACCENT
    pr2 = pat_p.add_run(
        "- https://openai.com/ko-KR/index/introducing-chatgpt-images-2-0/ ➔ public/gpt-image-2.jpeg (공식 4대 기능 맥락 반영)\n"
        "- https://www.analyticsvidhya.com/... ➔ public/GPT-image-2.webp (리더보드 1위 점수 및 벤치마크 분석 반영)\n"
        "👉 효과: 에이전트가 웹페이지를 직접 읽어 원문 팩트와 수치를 100% 정확하게 슬라이드에 반영!"
    )
    pr2.font.name = "Pretendard"
    pr2.font.size = Pt(9.5)
    pr2.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ─────────────────────────────────────────────────────────
    # SECTION 4: Phase 3 Deep-Dive (talks-slidev & 6 Standard Layouts)
    # ─────────────────────────────────────────────────────────
    sec4 = doc.add_paragraph()
    sec4_run = sec4.add_run("4. Phase 3: Slidev 덱 프로덕션 6대 표준 레이아웃")
    sec4_run.font.name = "Pretendard"
    sec4_run.font.size = Pt(14)
    sec4_run.font.bold = True
    sec4_run.font.color.rgb = PRIMARY
    sec4.paragraph_format.space_before = Pt(14)
    sec4.paragraph_format.space_after = Pt(6)

    layout_p = doc.add_paragraph()
    layout_p.paragraph_format.space_after = Pt(6)
    add_bullet_point(layout_p, "1. CoverSlide.vue (커버)", "WebGL Dynamic Fluid 캔버스 + 순백의 드롭섀도우 타이틀 + 동일 규격(h-6.5 px-2.5 py-1) 로고 바운딩 박스")
    add_bullet_point(layout_p, "2. SectionPartDivider.vue (디바이더)", "54:46 2열 분할 (좌측: 타이틀 & 핸드드로잉 브러시 언더라인 / 우측: 16:10 미디어 퓨어 프레임)")
    add_bullet_point(layout_p, "3. LiquidGlass 다열 카드 그리드 (개념)", "GPU 가속 순수 CSS 글래스모피즘 + 상하 여백 균형 + 1줄 캘리브레이션 푸터 태그")
    add_bullet_point(layout_p, "4. 인터랙티브 클릭 스테이지 (단계)", "useSlideContext() 기반 전환 + <div v-click=\"1\" class=\"hidden\"></div> AST 앵커 자동 배치")
    add_bullet_point(layout_p, "5. 다이렉트 에셋 쇼케이스 (에셋)", "이중 껍데기 박스 없이 h-82, h-92 상하 공간을 꽉 채운 1:1, 16:9, 9:16 다이렉트 렌더링")
    add_bullet_point(layout_p, "6. 실전 프롬프트 템플릿 & 실무 미션", "교육생 복사용 Case 1(Chat) & Case 2(Work) 카드 + 고정 시간(분) 없는 유연한 실습 카드")

    # ─────────────────────────────────────────────────────────
    # SECTION 5: The 6 Golden Guardrails
    # ─────────────────────────────────────────────────────────
    sec5 = doc.add_paragraph()
    sec5_run = sec5.add_run("5. 6대 무결점 품질 가드레일 (The 6 Golden Rules)")
    sec5_run.font.name = "Pretendard"
    sec5_run.font.size = Pt(14)
    sec5_run.font.bold = True
    sec5_run.font.color.rgb = PRIMARY
    sec5.paragraph_format.space_before = Pt(14)
    sec5.paragraph_format.space_after = Pt(6)

    # 6 Rules Table
    r_table = doc.add_table(rows=7, cols=3)
    r_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_table.autofit = False
    r_widths = [Inches(1.8), Inches(3.2), Inches(1.8)]
    
    r_headers = ["가드레일 영역", "절대 준수 규칙 (Best Practice)", "금지 사항 (Don'ts)"]
    for i, h in enumerate(r_headers):
        c = r_table.cell(0, i)
        c.width = r_widths[i]
        set_cell_shading(c, "0E4C92")
        set_cell_margins(c, 100, 100, 120, 120)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Pretendard"
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    rules_data = [
        ("① 한국어 타이포그래피", "word-break: keep-all 전역 적용\n카드 푸터/좁은 열 whitespace-nowrap\n상하 여백 휑하지 않게 패딩 캘리브레이션", "마지막 1글자만 줄바꿈되는 현상\n(Orphan Wrap: ~화, ~결, ~장)"),
        ("② 에셋 순수성", "퍼블릭 에셋/표/UI 원본 다이렉트 렌더링\nh-82, h-92로 상하 여백 꽉 채우기\n로고 바운딩 박스 규격 1:1 통일", "이미지에 불필요한 <LiquidGlass>나\n두꺼운 보더 박스(이중 껍데기) 씌우기"),
        ("③ 120 FPS 렌더링 성능", "GPU 가속 순수 CSS 글래스모피즘\nslides.md에 preload: true (0ms 탐색)\nglobal-bottom에 contain: strict 격리", "50장 이상 덱에서 무거운\nSVG Displacement Filter (Refractive) 사용"),
        ("④ 클릭 스텝 정합성", "Custom 컴포넌트 내 useSlideContext 연동\n<div v-click=\"1\" class=\"hidden\"></div> 앵커", "동작하지 않는 빈 클릭(Dead Click)\n(단일 뷰는 과감히 clicks: 0 적용)"),
        ("⑤ 강의 기획 & Anti-AI", "실무 미션 카드(실습 과제: ...)로 명시\n모든 슬라이드 하단에 구어체 [강사 멘트]\n즉시 복사 가능한 프롬프트 치트키", "실습 슬라이드에 고정 시간(25분, 35분)\n라이브 데모 전용 슬라이드 (강사 직접시연)\n상투적인 AI 미사여구 및 난해한 기술용어"),
        ("⑥ 안전한 로컬 Git 관리", "빌드 통과(built in 3~4s) 확인 후\ngit commit -m \"...\" 로컬 커밋만 수행", "원격 저장소(origin)에 git push 실행\n(절대 원격 푸시 금지)")
    ]

    for row_idx, data in enumerate(rules_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            c = r_table.cell(row_idx, col_idx)
            c.width = r_widths[col_idx]
            set_cell_shading(c, bg)
            set_cell_margins(c, 100, 100, 120, 120)
            set_cell_border(c, bottom={"val": "single", "sz": "4", "color": "E2E8F0"})
            p = c.paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = "Pretendard"
            r.font.size = Pt(9)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = PRIMARY
            elif col_idx == 2:
                r.font.color.rgb = RGBColor(225, 29, 72) # Rose 600
            else:
                r.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─────────────────────────────────────────────────────────
    # SECTION 6: Cheat Sheet
    # ─────────────────────────────────────────────────────────
    sec6 = doc.add_paragraph()
    sec6_run = sec6.add_run("6. 강사용 실전 명령어 & 프롬프트 치트키 (Quick Reference)")
    sec6_run.font.name = "Pretendard"
    sec6_run.font.size = Pt(14)
    sec6_run.font.bold = True
    sec6_run.font.color.rgb = PRIMARY
    sec6.paragraph_format.space_before = Pt(14)
    sec6.paragraph_format.space_after = Pt(6)

    # Code Box Table
    t_code = doc.add_table(rows=1, cols=1)
    t_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_code = t_code.cell(0, 0)
    c_code.width = Inches(6.8)
    set_cell_shading(c_code, "0F172A") # Slate 900
    set_cell_margins(c_code, 140, 140, 160, 160)
    
    code_p = c_code.paragraphs[0]
    code_p.paragraph_format.space_after = Pt(0)
    code_text = (
        "# 1. 기획 단계 (lecture-director 호출)\n"
        "\"이 교육 계획서와 유튜브 링크 참고해서 lecture-director 스킬로 브리프 짜줘.\"\n\n"
        "# 2. 덱 제작 단계 (talks-slidev 호출)\n"
        "\"sessions/ 브리프와 public/ 에셋들 기반으로 talks-slidev 스킬로 슬라이드 완성해줘.\"\n\n"
        "# 3. 빌드 무결점 검증 & 로컬 커밋\n"
        "pnpm run build\n"
        "git add -A; git commit -m \"feat: build Talks-grade slide deck\"\n\n"
        "# 4. 교육생 실전 프롬프트 4단 공식\n"
        "[목적/배경] ➔ [화면비/2K] ➔ [\"큰따옴표 한글 인쇄 문구\"] ➔ [포토리얼리즘 스타일]"
    )
    c_run = code_p.add_run(code_text)
    c_run.font.name = "Fira Code"
    c_run.font.size = Pt(9)
    c_run.font.color.rgb = RGBColor(56, 189, 248) # Light Cyan

    # Save DOCX
    output_path = r"C:\Users\IN\.gemini\antigravity\scratch\ai-lecture-curriculum\260908_mc-energy\lecture_slide_production_master_guide.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_master_guide()
